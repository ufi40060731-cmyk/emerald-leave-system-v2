from __future__ import annotations

import csv
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable


SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".json", ".csv", ".pdf", ".docx"}


@dataclass(frozen=True)
class LoadedDocument:
    source: str
    title: str
    text: str


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8-sig", errors="replace")


def _load_json(path: Path) -> str:
    payload = json.loads(_read_text(path))
    return json.dumps(payload, ensure_ascii=False, indent=2)


def _load_csv(path: Path) -> str:
    rows: list[str] = []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        for row in csv.reader(handle):
            rows.append(" | ".join(cell.strip() for cell in row))
    return "\n".join(rows)


def _load_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise RuntimeError("PDF support requires: pip install pypdf") from exc
    reader = PdfReader(str(path))
    return "\n\n".join((page.extract_text() or "").strip() for page in reader.pages)


def _load_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise RuntimeError("DOCX support requires: pip install python-docx") from exc
    document = Document(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def load_document(path: Path, root: Path | None = None) -> LoadedDocument:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported document type: {path.suffix}")

    if suffix in {".txt", ".md", ".markdown"}:
        text = _read_text(path)
    elif suffix == ".json":
        text = _load_json(path)
    elif suffix == ".csv":
        text = _load_csv(path)
    elif suffix == ".pdf":
        text = _load_pdf(path)
    else:
        text = _load_docx(path)

    source = path.relative_to(root).as_posix() if root else path.as_posix()
    return LoadedDocument(source=source, title=path.stem.replace("_", " "), text=text.strip())


def discover_documents(root: Path) -> Iterable[Path]:
    for path in sorted(root.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            yield path
